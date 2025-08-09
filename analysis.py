# analysis.py
import pandas as pd, numpy as np
import scipy.stats as stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.anova import anova_lm
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.anova import AnovaRM
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path
import io

def load_data(path):
    df = pd.read_csv(path)
    return df

def one_way_anova(df, factor, value):
    formula = f'{value} ~ C({factor})'
    model = ols(formula, data=df).fit()
    anova = anova_lm(model, typ=2)
    return model, anova

def two_way_anova(df, factor_a, factor_b, value, interaction=True):
    if interaction:
        formula = f'{value} ~ C({factor_a}) * C({factor_b})'
    else:
        formula = f'{value} ~ C({factor_a}) + C({factor_b})'
    model = ols(formula, data=df).fit()
    anova = anova_lm(model, typ=2)
    return model, anova

def three_way_anova(df, a, b, c, value, interaction=True):
    if interaction:
        formula = f'{value} ~ C({a}) * C({b}) * C({c})'
    else:
        formula = f'{value} ~ C({a}) + C({b}) + C({c})'
    model = ols(formula, data=df).fit()
    anova = anova_lm(model, typ=2)
    return model, anova

def repeated_measures_anova(df, subject_col, within_col, dv_col):
    # df must contain subject identifier and within factor (within_col can be list)
    if isinstance(within_col, list):
        aovrm = AnovaRM(df, dv_col, subject_col, within=within_col).fit()
    else:
        aovrm = AnovaRM(df, dv_col, subject_col, within=[within_col]).fit()
    return aovrm

def shapiro_test(series):
    series = series.dropna()
    if len(series) < 3:
        return (None, None)
    w, p = stats.shapiro(series)
    return w, p

def levene_test(df, group_col, value_col):
    groups = [g[value_col].values for n,g in df.groupby(group_col)]
    if len(groups) < 2:
        return (None, None)
    stat, p = stats.levene(*groups)
    return stat, p

def tukey_hsd(df, group_col, value_col):
    res = pairwise_tukeyhsd(df[value_col], df[group_col])
    return res

def plot_box(df, group_col, value_col, out_path):
    fig, ax = plt.subplots(figsize=(6,4))
    df.boxplot(column=value_col, by=group_col, ax=ax)
    ax.set_title(f'Boxplot by {group_col}')
    ax.set_ylabel(value_val := value_value if False else value_col)
    plt.suptitle("")
    fig.savefig(out_path, bbox_inches='tight')
    plt.close(fig)
    return out_path

def plot_hist(series, out_path):
    fig, ax = plt.subplots(figsize=(6,4))
    ax.hist(series, bins=12)
    ax.set_title('Histogram')
    fig.savefig(out_path, bbox_inches='tight')
    plt.close(fig)
    return out_path

def plot_interaction(df, factor_a, factor_b, value_col, out_path):
    means = df.groupby([factor_a, factor_b])[value_col].mean().unstack()
    fig, ax = plt.subplots(figsize=(6,4))
    means.T.plot(marker='o', ax=ax)
    ax.set_title(f'Interaction plot: {factor_b} by {factor_a}')
    ax.set_xlabel(factor_b)
    ax.set_ylabel(value_col)
    fig.savefig(out_path, bbox_inches='tight')
    plt.close(fig)
    return out_path

def generate_report(docx_path, df, analyses_results, plots):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("Звіт: АНOVA аналіз (Статистика)")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph("Вихідні дані (зразок):")
    tbl = doc.add_table(rows=1, cols= len(df.columns))
    hdr = tbl.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, r in df.head(30).iterrows():
        row = tbl.add_row().cells
        for i, c in enumerate(df.columns):
            row[i].text = str(r[c])
    doc.add_paragraph()
    # Insert analyses text
    for k,v in analyses_results.items():
        doc.add_paragraph(k + ":")
        doc.add_paragraph(str(v))
    # Insert plots
    for ppath in plots:
        if Path(ppath).exists():
            doc.add_picture(ppath, width=Inches(5))
            doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = 1
    para.add_run('\n\nРозробник: Чаплоуцький А.М., кафедра плодівництва і виноградарства, УНУ')
    doc.save(docx_path)
    return docx_path
